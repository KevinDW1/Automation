"""
run_all_qa.py
=============
Master runner -- executes all QA suites, captures results,
then regenerates QA_Validation_Report.docx automatically.

Usage
-----
  python run_all_qa.py              # run all suites + generate report
  python run_all_qa.py --report     # regenerate report from last results only
  python run_all_qa.py --suite INV  # run one suite + update report
  python run_all_qa.py --suite JOB
  python run_all_qa.py --suite VEN
  python run_all_qa.py --suite EQP
"""

from __future__ import annotations
import json, os, re, subprocess, sys, time
from datetime import datetime
from pathlib import Path

HERE         = Path(__file__).parent
RESULTS_FILE = HERE / "qa_results.json"
REPORT_FILE  = HERE / "QA_Validation_Report.docx"

SUITES = {
    "INV": {"label": "Invoice Suite",    "script": HERE / "Invoice"   / "invoice_suite.py",    "prefix": "INV-"},
    "JOB": {"label": "Jobsite Suite",    "script": HERE / "Jobsite"   / "jobsite_suite.py",    "prefix": "JOB-"},
    "VEN": {"label": "Vendor Suite",     "script": HERE / "Vendor"    / "vendor_suite.py",     "prefix": "VEN-"},
    "CUS": {"label": "Customer Suite",   "script": HERE / "Customer"  / "cus02_to_cus24.py",   "prefix": "CUS-"},
    "EQP": {"label": "Equipment Suite",  "script": HERE / "Equipment-Management" / "tc01_to_tc22.py", "prefix": "TC-"},
}

def parse_output(output: str, prefix: str) -> list[dict]:
    results = []
    pattern = re.compile(
        r"={40,}\n\s*(" + re.escape(prefix) + r"\d+)\s+--\s+(.+?)\n"
        r"-{40,}\n\s*VERDICT:\s*(PASS|FAIL)(.*?)\n={40,}",
        re.DOTALL
    )
    for m in pattern.finditer(output):
        body     = m.group(4)
        failures = re.findall(r"\*\s+(.+)", body)
        evidence = [e.strip() for e in re.findall(r"^\s{4}(.+)$", body, re.MULTILINE)
                    if e.strip() and not e.strip().startswith("*")]
        results.append({
            "id":       m.group(1).strip(),
            "title":    m.group(2).strip(),
            "passed":   m.group(3).strip() == "PASS",
            "failures": failures,
            "evidence": evidence[:4],
        })
    return results

def run_suite(key: str) -> list[dict]:
    suite  = SUITES[key]
    script = Path(suite["script"])
    if not script.exists():
        print(f"  SKIP {key} -- {script} not found")
        return []
    print(f"\n{'='*60}\n  Running {suite['label']}...\n{'='*60}")
    start = time.time()
    proc  = subprocess.run(
        [sys.executable, str(script)],
        capture_output=True, text=True, cwd=str(script.parent),
    )
    elapsed = round(time.time() - start, 1)
    output  = proc.stdout + proc.stderr
    results = parse_output(output, suite["prefix"])
    passed  = sum(1 for r in results if r["passed"])
    failed  = len(results) - passed
    print(f"  {suite['label']} done in {elapsed}s -- {passed} PASS / {failed} FAIL")
    if not results:
        print("  WARNING: No results parsed -- check script output format")
    return results

def load_results() -> dict:
    if RESULTS_FILE.exists():
        try:
            return json.loads(RESULTS_FILE.read_text(encoding="utf-8"))
        except Exception:
            pass
    return {}

def save_results(data: dict) -> None:
    RESULTS_FILE.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")

def generate_report(all_results: dict) -> None:
    print(f"\n{'='*60}\n  Generating report...\n{'='*60}")
    # Build report data
    sections = []
    for key, data in all_results.items():
        if key.startswith("_"):
            continue
        res    = data.get("results", [])
        passed = sum(1 for r in res if r.get("passed"))
        failed = sum(1 for r in res if not r.get("passed"))
        sections.append({
            "key":    key,
            "label":  data.get("label", key),
            "run_at": data.get("run_at", "Not run"),
            "total":  len(res),
            "passed": passed,
            "failed": failed,
            "results": res,
        })

    # Write plain-text report (works without node/docx)
    now   = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    lines = []
    lines.append("=" * 70)
    lines.append("  WASTE APPLICATIONS -- QA VALIDATION REPORT")
    lines.append(f"  Generated: {now}")
    lines.append("=" * 70)
    lines.append("")

    # Summary
    lines.append("EXECUTIVE SUMMARY")
    lines.append("-" * 70)
    lines.append(f"  {'Suite':<25} {'Total':>6} {'Pass':>6} {'Fail':>6} {'Pct':>6}  {'Last Run'}")
    lines.append(f"  {'-'*25} {'-'*6} {'-'*6} {'-'*6} {'-'*6}  {'-'*20}")
    grand_total = grand_pass = grand_fail = 0
    for s in sections:
        pct = f"{round(s['passed']/s['total']*100)}%" if s['total'] > 0 else "N/A"
        lines.append(f"  {s['label']:<25} {s['total']:>6} {s['passed']:>6} {s['failed']:>6} {pct:>6}  {s['run_at']}")
        grand_total += s["total"]
        grand_pass  += s["passed"]
        grand_fail  += s["failed"]
    lines.append(f"  {'-'*25} {'-'*6} {'-'*6} {'-'*6} {'-'*6}")
    grand_pct = f"{round(grand_pass/grand_total*100)}%" if grand_total > 0 else "N/A"
    lines.append(f"  {'TOTAL':<25} {grand_total:>6} {grand_pass:>6} {grand_fail:>6} {grand_pct:>6}")
    lines.append("")

    # Detail per suite
    for s in sections:
        lines.append("")
        lines.append("=" * 70)
        lines.append(f"  {s['label'].upper()}  --  {s['passed']} PASS / {s['failed']} FAIL  --  {s['run_at']}")
        lines.append("=" * 70)
        lines.append(f"  {'ID':<10} {'Result':<8} {'Title'}")
        lines.append(f"  {'-'*10} {'-'*8} {'-'*48}")
        if not s["results"]:
            lines.append("  (Pending -- suite not yet executed)")
        for r in s["results"]:
            verdict = "PASS" if r.get("passed") else "FAIL"
            lines.append(f"  {r['id']:<10} {verdict:<8} {r['title']}")
            for f in r.get("failures", []):
                lines.append(f"  {'':10} {'':8}   * {f}")

    lines.append("")
    lines.append("=" * 70)
    lines.append("  END OF REPORT")
    lines.append("=" * 70)

    report_txt = HERE / "QA_Validation_Report.txt"
    report_txt.write_text("\n".join(lines), encoding="utf-8")
    print(f"  OK Text report: {report_txt}")

    # Also try to generate docx if node is available
    try:
        generate_docx(all_results, sections, now)
    except Exception as e:
        print(f"  DOCX generation skipped: {e}")
        print(f"  Text report available at: {report_txt}")

def generate_docx(all_results: dict, sections: list, now: str) -> None:
    """Generate Word docx using Python python-docx."""
    try:
        from docx import Document as DocxDoc
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise Exception("python-docx not installed -- run: pip install python-docx")

    doc = DocxDoc()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(0.75)
        section.right_margin  = Inches(0.75)

    # Helper: set cell bg colour
    def set_cell_bg(cell_obj, hex_color):
        tc   = cell_obj._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        tcPr.append(shd)

    def rgb(hex_str):
        h = hex_str.lstrip("#")
        return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("QA Validation Report")
    run.bold      = True
    run.font.size = Pt(24)
    run.font.color.rgb = rgb(NAVY)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r2 = sub.add_run(f"Waste Applications  |  {now}")
    r2.font.size = Pt(11)
    r2.font.color.rgb = rgb("444444")

    doc.add_paragraph()

    # ── Summary table ────────────────────────────────────────────────────────
    doc.add_heading("Executive Summary", level=1)
    hdrs  = ["Suite", "Total", "Pass", "Fail", "Pass %", "Last Run"]
    widths = [4, 1.2, 1.2, 1.2, 1.2, 3]
    tbl = doc.add_table(rows=1, cols=len(hdrs))
    tbl.style = "Table Grid"
    hdr_row = tbl.rows[0]
    for i, (h, w) in enumerate(zip(hdrs, widths)):
        c = hdr_row.cells[i]
        set_cell_bg(c, NAVY)
        c.width = Inches(w / 2)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = rgb(WHITE)
        run.font.size = Pt(9)

    for s in sections:
        pct  = f"{round(s['passed']/s['total']*100)}%" if s['total'] > 0 else "N/A"
        vals = [s["label"], str(s["total"]), str(s["passed"]), str(s["failed"]), pct, s["run_at"]]
        row  = tbl.add_row()
        for i, (v, w) in enumerate(zip(vals, widths)):
            c = row.cells[i]
            c.width = Inches(w / 2)
            fill = WHITE
            col  = "000000"
            if i == 2 and s["passed"] > 0:   fill = "DFF2E0"; col = GREEN
            if i == 3 and s["failed"] > 0:   fill = "FDECEA"; col = RED
            if i == 3 and s["failed"] == 0:  fill = WHITE
            set_cell_bg(c, fill)
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(v)
            r.font.size = Pt(9)
            r.font.color.rgb = rgb(col)
            if i == 3 and s["failed"] > 0: r.bold = True
            if i == 2 and s["passed"] > 0: r.bold = True

    # Totals row
    grand_t = sum(s["total"]  for s in sections)
    grand_p = sum(s["passed"] for s in sections)
    grand_f = sum(s["failed"] for s in sections)
    grand_pct = f"{round(grand_p/grand_t*100)}%" if grand_t > 0 else "N/A"
    tot_row = tbl.add_row()
    for i, v in enumerate(["TOTAL", str(grand_t), str(grand_p), str(grand_f), grand_pct, ""]):
        c = tot_row.cells[i]
        set_cell_bg(c, NAVY)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(v)
        r.bold = True
        r.font.color.rgb = rgb(WHITE)
        r.font.size = Pt(9)

    doc.add_paragraph()

    # ── Per-suite detail ─────────────────────────────────────────────────────
    NAVY_STR  = NAVY
    GREEN_STR = GREEN
    RED_STR   = RED
    AMBER_STR = AMBER

    for s in sections:
        doc.add_page_break()
        doc.add_heading(f"{s['label']}", level=1)
        info = doc.add_paragraph()
        r = info.add_run(
            f"{s['passed']} PASS  /  {s['failed']} FAIL"
            + (f"  |  Last run: {s['run_at']}" if s['run_at'] != 'Not run' else "  |  Not yet executed")
        )
        r.font.size = Pt(10)
        r.font.color.rgb = rgb("444444")
        doc.add_paragraph()

        if not s["results"]:
            p = doc.add_paragraph()
            r = p.add_run("Suite pending — not yet executed.")
            r.font.color.rgb = rgb(AMBER_STR)
            r.italic = True
            continue

        tbl2 = doc.add_table(rows=1, cols=4)
        tbl2.style = "Table Grid"
        col_w = [0.8, 2.8, 3.8, 1.0]
        h_row = tbl2.rows[0]
        for i, (h, w) in enumerate(zip(["ID", "Scenario", "Detail / Evidence", "Result"], col_w)):
            c = h_row.cells[i]
            set_cell_bg(c, NAVY_STR)
            c.width = Inches(w)
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(h)
            run.bold = True
            run.font.color.rgb = rgb(WHITE)
            run.font.size = Pt(9)

        for idx, r in enumerate(s["results"]):
            row = tbl2.add_row()
            fill = "F5F5F5" if idx % 2 == 0 else WHITE

            # ID
            c = row.cells[0]; set_cell_bg(c, fill); c.width = Inches(col_w[0])
            p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rn = p.add_run(r["id"]); rn.bold = True; rn.font.size = Pt(9)

            # Scenario / title
            c = row.cells[1]; set_cell_bg(c, fill); c.width = Inches(col_w[1])
            p = c.paragraphs[0]
            rn = p.add_run(r["title"]); rn.font.size = Pt(9)

            # Detail
            c = row.cells[2]; set_cell_bg(c, fill); c.width = Inches(col_w[2])
            p = c.paragraphs[0]
            detail = "; ".join(r.get("failures", []) + r.get("evidence", []))
            if not detail:
                detail = "See evidence log for details."
            rn = p.add_run(detail[:200]); rn.font.size = Pt(8); rn.font.color.rgb = rgb("333333")

            # Verdict
            passed  = r.get("passed")
            vfill   = "DFF2E0" if passed else "FDECEA"
            vtxt    = "PASS" if passed else "FAIL"
            vcol    = GREEN_STR if passed else RED_STR
            c = row.cells[3]; set_cell_bg(c, vfill); c.width = Inches(col_w[3])
            p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rn = p.add_run(vtxt); rn.bold = True; rn.font.size = Pt(9)
            rn.font.color.rgb = rgb(vcol)

        doc.add_paragraph()

    doc.save(str(REPORT_FILE))
    print(f"  OK DOCX report: {REPORT_FILE} ({REPORT_FILE.stat().st_size:,} bytes)")

NAVY  = "1B3A5C"
TEAL  = "0D6E6E"
GREEN = "1D6A2F"
RED   = "8B1A1A"
AMBER = "7A4F00"
WHITE = "FFFFFF"

def main() -> None:
    args = sys.argv[1:]

    if "--report" in args:
        all_results = load_results()
        if not all_results:
            print("No results found. Run suites first.")
            sys.exit(1)
        sections = build_sections(all_results)
        generate_report(all_results)
        return

    suite_filter = None
    if "--suite" in args:
        idx = args.index("--suite")
        if idx + 1 < len(args):
            suite_filter = args[idx + 1].upper()
            if suite_filter not in SUITES:
                print(f"Unknown suite: {suite_filter}. Valid: {', '.join(SUITES)}")
                sys.exit(1)

    all_results = load_results()
    all_results["_meta"] = {
        "last_run":   datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "suites_run": suite_filter or "ALL",
    }

    suites_to_run = [suite_filter] if suite_filter else list(SUITES.keys())
    for key in suites_to_run:
        results = run_suite(key)
        all_results[key] = {
            "label":   SUITES[key]["label"],
            "run_at":  datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "results": results,
        }

    save_results(all_results)
    print(f"\n  Results saved to {RESULTS_FILE.name}")
    generate_report(all_results)

    # Summary
    print(f"\n{'='*60}\n  FINAL SUMMARY\n{'='*60}")
    for key, data in all_results.items():
        if key.startswith("_"):
            continue
        res    = data.get("results", [])
        passed = sum(1 for r in res if r.get("passed"))
        failed = len(res) - passed
        pct    = f"{round(passed/len(res)*100)}%" if res else "N/A"
        print(f"  {key:4}  {passed:2} PASS  {failed:2} FAIL  {pct:>5}  {data.get('run_at','')}")
    print(f"{'='*60}")
    print(f"\n  Report: QA_Validation_Report.docx / .txt")

def build_sections(all_results):
    sections = []
    for key, data in all_results.items():
        if key.startswith("_"): continue
        res = data.get("results", [])
        sections.append({
            "key": key, "label": data.get("label", key),
            "run_at": data.get("run_at", "Not run"),
            "total": len(res),
            "passed": sum(1 for r in res if r.get("passed")),
            "failed": sum(1 for r in res if not r.get("passed")),
            "results": res,
        })
    return sections

if __name__ == "__main__":
    main()
